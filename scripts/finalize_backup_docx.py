from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document


ROOT = Path(r"C:\Projects\NutritionApp")
FINAL_PATH = ROOT / "Диплом Абдулин Дмитрий АИС-42 - backup.docx"
WORK_PATH = ROOT / "Диплом Абдулин Дмитрий АИС-42 - backup.updated.docx"


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def append_sentence(paragraph, text: str) -> None:
    base = paragraph.text.rstrip()
    if not base:
        new_text = text
    elif base.endswith((".", "!", "?", ":", ";")):
        new_text = f"{base} {text}".strip()
    else:
        new_text = f"{base} {text}".strip()
    set_paragraph_text(paragraph, new_text)


def set_style(doc: Document, idx: int, style_name: str) -> None:
    doc.paragraphs[idx].style = doc.styles[style_name]


def replace_references(doc: Document) -> None:
    references = {
        855: "Федеральный закон от 27.07.2006 № 149-ФЗ «Об информации, информационных технологиях и о защите информации».",
        856: "Федеральный закон от 27.07.2006 № 152-ФЗ «О персональных данных».",
        857: "ГОСТ 34.601–90. Информационная технология. Комплекс стандартов на автоматизированные системы. Автоматизированные системы. Стадии создания.",
        858: "МР 2.3.1.0253-21. Нормы физиологических потребностей в энергии и пищевых веществах для различных групп населения Российской Федерации [Электронный ресурс]. – Режим доступа: https://internet-law.ru/documents/dop_documents/10/gost_35215/0/mr_2.html (дата обращения: 23.05.2026).",
        859: "Методические указания по выполнению выпускной квалификационной работы для обучающихся по направлению 09.03.03 «Прикладная информатика» / Уральский государственный лесотехнический университет. Кафедра интеллектуальных систем. – Екатеринбург, 2025. – 42 с.",
        860: "Барсегян А. А., Куприянов М. С., Степаненко В. В., Холод И. И. Технологии анализа данных: Data Mining, Visual Mining, Text Mining, OLAP. – СПб.: БХВ-Петербург, 2021. – 384 с.",
        861: "Баланов А. Н. Автоматизация, цифровизация и оптимизация бизнес-процессов: IT-решения и стратегии для современных компаний. – СПб.: Лань, 2025. – 48 с.",
        862: "Разработка рекомендательной системы на основе сессий с использованием многоуровневой системы отбора кандидатов // КиберЛенинка [Электронный ресурс]. – Режим доступа: https://cyberleninka.ru/article/n/razrabotka-rekomendatelnoy-sistemy-na-osnove-sessiy-s-ispolzovaniem-mnogourovnevoy-sistemy-otbora-kandidatov (дата обращения: 23.05.2026).",
        863: "Обзор методов построения рекомендательных систем на основе сессий // КиберЛенинка [Электронный ресурс]. – Режим доступа: https://cyberleninka.ru/article/n/obzor-metodov-postroeniya-rekomendatelnyh-sistem-na-osnove-sessiy (дата обращения: 23.05.2026).",
        864: "Контент-ориентированный подход в системах рекомендаций: принципы, методы и метрики эффективности // КиберЛенинка [Электронный ресурс]. – Режим доступа: https://cyberleninka.ru/article/n/kontent-orientirovannyy-podhod-v-sistemah-rekomendatsiy-printsipy-metody-i-metriki-effektivnosti (дата обращения: 23.05.2026).",
        865: "Разработка алгоритмов ранжирования с контекстной адаптацией для рекомендательных систем // КиберЛенинка [Электронный ресурс]. – Режим доступа: https://cyberleninka.ru/article/n/razrabotka-algoritmov-ranzhirovaniya-s-kontekstnoy-adaptatsiey-dlya-rekomendatelnyh-sistem (дата обращения: 23.05.2026).",
        866: "Машинное обучение модели информационной рекомендательной системы по вопросам индивидуализации образования // КиберЛенинка [Электронный ресурс]. – Режим доступа: https://cyberleninka.ru/article/n/mashinnoe-obuchenie-modeli-informatsionnoy-rekomendatelnoy-sistemy-po-voprosam-individualizatsii-obrazovaniya (дата обращения: 23.05.2026).",
        867: "Android Developers. Build better apps with Jetpack Compose [Электронный ресурс]. – Режим доступа: https://developer.android.com/jetpack/compose (дата обращения: 23.05.2026).",
        868: "Android Developers. Guide to app architecture [Электронный ресурс]. – Режим доступа: https://developer.android.com/topic/architecture (дата обращения: 23.05.2026).",
        869: "Spring. Spring Boot Reference Documentation [Электронный ресурс]. – Режим доступа: https://docs.spring.io/spring-boot/documentation.html (дата обращения: 23.05.2026).",
        870: "PostgreSQL Global Development Group. PostgreSQL Documentation [Электронный ресурс]. – Режим доступа: https://www.postgresql.org/docs/ (дата обращения: 23.05.2026).",
        871: "Google Developers. ML Kit for Android [Электронный ресурс]. – Режим доступа: https://developers.google.com/ml-kit (дата обращения: 23.05.2026).",
        872: "Jsoup. Java HTML Parser. Official Documentation [Электронный ресурс]. – Режим доступа: https://jsoup.org/ (дата обращения: 23.05.2026).",
        873: "Tess4J. Java JNA wrapper for Tesseract OCR [Электронный ресурс]. – Режим доступа: https://tess4j.sourceforge.net/ (дата обращения: 23.05.2026).",
        874: "OWASP Foundation. OWASP Top 10 Web Application Security Risks [Электронный ресурс]. – Режим доступа: https://owasp.org/www-project-top-ten/ (дата обращения: 23.05.2026).",
    }
    for idx, text in references.items():
        set_paragraph_text(doc.paragraphs[idx], text)


def fix_paragraphs_and_captions(doc: Document) -> None:
    # Normalize body styles that visually drift in Word.
    for idx in range(114, 855):
        paragraph = doc.paragraphs[idx]
        if paragraph.style.name in {"Normal (Web)", "break-words"}:
            set_style(doc, idx, "Normal")

    for idx in [470, 597, 598, 599, 600, 601]:
        set_style(doc, idx, "Normal")

    # Restore table titles and captions using the original template positions.
    append_sentence(doc.paragraphs[164], "анализ истории питания и формирование рекомендаций.")
    set_paragraph_text(doc.paragraphs[165], "Таблица 1.1 – Основные характеристики текущей версии системы")
    set_style(doc, 165, "Normal")

    append_sentence(doc.paragraphs[170], "Именно это делает задачу автоматизации планирования рациона практически значимой и технологически обоснованной.")
    set_paragraph_text(doc.paragraphs[171], "Рисунок 1.1 – Контекст взаимодействия пользователя и системы")
    set_style(doc, 171, "Caption")

    append_sentence(doc.paragraphs[174], "Мобильный клиент обеспечивает интерфейс ввода, локальное кеширование данных и обращение к REST API.")
    set_paragraph_text(doc.paragraphs[176], "Таблица 1.2 – Участники и роли в процессе эксплуатации системы")
    set_style(doc, 176, "Normal")

    append_sentence(doc.paragraphs[186], "Такой подход не обеспечивает целостности данных и плохо масштабируется на длительный период наблюдения.")
    set_paragraph_text(doc.paragraphs[187], "Таблица 1.3 – Недостатки ручного ведения питания и их последствия")
    set_style(doc, 187, "Normal")

    append_sentence(doc.paragraphs[255], "Отдельные продукты обладают мощной аналитикой, но требуют подписки уже на базовых сценариях персонализации.")
    set_paragraph_text(doc.paragraphs[256], "Таблица 1.4 – Сравнительный анализ существующих цифровых решений")
    set_style(doc, 256, "Normal")

    append_sentence(doc.paragraphs[298], "На этапе проектирования определены архитектурные границы между мобильным клиентом, API и базой данных.")
    set_paragraph_text(doc.paragraphs[299], "Таблица 2.1 – Этапы разработки системы")
    set_style(doc, 299, "Normal")

    append_sentence(doc.paragraphs[347], "Рецепт представляет собой агрегат, содержащий описание, инструкцию, время приготовления, КБЖУ и список ингредиентов.")
    set_paragraph_text(doc.paragraphs[348], "Рисунок 2.1 – Упрощенная ER-модель серверной базы данных")
    set_style(doc, 348, "Caption")

    append_sentence(doc.paragraphs[379], "Ключевые объекты системы и их назначение приведены в таблице 2.3.")
    set_paragraph_text(doc.paragraphs[380], "Таблица 2.3 – Ключевые объекты данных системы")
    set_style(doc, 380, "Normal")

    append_sentence(doc.paragraphs[394], "Структура аналитического блока включает:")
    set_paragraph_text(doc.paragraphs[395], "Рисунок 2.3 – Экран аналитики и отчетных показателей")
    set_style(doc, 395, "Caption")

    set_paragraph_text(doc.paragraphs[405], "краткое объяснение причины рекомендации; карточки с изображением, названием и базовыми нутриционными характеристиками;")
    set_paragraph_text(doc.paragraphs[406], "Рисунок 2.4 – Экран каталога и рекомендаций рецептов")
    set_style(doc, 406, "Caption")

    append_sentence(doc.paragraphs[458], "Рекомендательная подсистема имеет гибридную структуру и реализована на сервере вокруг сервиса PersonalizedRecommendationService и модуля ML-ранжирования.")
    set_paragraph_text(doc.paragraphs[460], "Рисунок 2.5 – Фрагмент реализации REST-интерфейса и клиентского вызова")
    set_style(doc, 460, "Caption")

    set_paragraph_text(doc.paragraphs[468], "Рисунок 2.6 – Логика гибридной рекомендательной подсистемы с ML-ранжированием")
    set_style(doc, 468, "Caption")

    append_sentence(doc.paragraphs[523], "Экран дневника содержит список приемов пищи за выбранную дату и краткую сводку по дню.")
    set_paragraph_text(doc.paragraphs[524], "Рисунок 2.8 – Экран дневника питания")
    set_style(doc, 524, "Caption")

    append_sentence(doc.paragraphs[539], "Пользователь может перейти к детальной карточке рецепта, открыть его описание, инструкцию и состав ингредиентов.")
    set_paragraph_text(doc.paragraphs[540], "Рисунок 2.10 – Экран каталога рецептов")
    set_style(doc, 540, "Caption")

    append_sentence(doc.paragraphs[541], "Экран добавления еды поддерживает три режима:")
    set_paragraph_text(doc.paragraphs[542], "Рисунок 2.11 – Экран добавления еды")
    set_style(doc, 542, "Caption")

    append_sentence(doc.paragraphs[563], "Подбор выполняется через общий каталог продуктов, а результаты синхронизируются с профилем.")
    set_paragraph_text(doc.paragraphs[564], "Рисунок 2.14 – Экран профиля и персональных настроек")
    set_style(doc, 564, "Caption")

    set_paragraph_text(doc.paragraphs[576], "Рисунок 2.15 – Экран пищевых предпочтений")
    set_style(doc, 576, "Caption")
    set_paragraph_text(doc.paragraphs[584], "Рисунок 2.16 – Экран списка покупок")
    set_style(doc, 584, "Caption")
    set_paragraph_text(doc.paragraphs[595], "Рисунок 2.17 – Экран избранного")
    set_style(doc, 595, "Caption")
    set_paragraph_text(doc.paragraphs[607], "Рисунок 2.12 – Экран плана питания")
    set_style(doc, 607, "Caption")
    set_paragraph_text(doc.paragraphs[617], "Рисунок 2.13 – Экран холодильника")
    set_style(doc, 617, "Caption")

    append_sentence(doc.paragraphs[724], "Особое значение имела проверка длинных цепочек действий, где ошибка на одном экране влияет на последующее поведение системы.")
    set_paragraph_text(doc.paragraphs[725], "Таблица 2.4 – Результаты тестирования системы")
    set_style(doc, 725, "Normal")

    set_paragraph_text(doc.paragraphs[762], "Таблица 3.1 – Трудоемкость этапов проекта")
    set_style(doc, 762, "Normal")
    set_paragraph_text(
        doc.paragraphs[763],
        "Среднемесячная заработная плата участников проекта приведена в таблице 3.2. Для расчета затрат на оплату труда используется стандартный подход с учетом количества человеко-дней и страховых отчислений."
    )
    set_paragraph_text(doc.paragraphs[764], "Таблица 3.2 – Среднемесячная заработная плата участников проекта, руб.")
    set_style(doc, 764, "Normal")

    append_sentence(doc.paragraphs[773], "Поэтому амортизационные и энергетические расходы учитываются укрупненно.")
    set_paragraph_text(doc.paragraphs[774], "Таблица 3.3 – Затраты на оплату труда с отчислениями, руб.")
    set_style(doc, 774, "Normal")

    set_paragraph_text(doc.paragraphs[805], "Таблица 3.4 – Смета затрат на проект")
    set_style(doc, 805, "Normal")

    append_sentence(doc.paragraphs[819], "Чистый экономический эффект за первый год определяется как разность между годовой экономией и суммарными затратами проекта:")
    set_paragraph_text(doc.paragraphs[820], "Таблица 3.5 – Годовая экономия от внедрения")
    set_style(doc, 820, "Normal")

    # Final wording polish.
    set_paragraph_text(
        doc.paragraphs[852],
        "Перспективы дальнейшего развития системы связаны с расширением набора признаков, учетом сезонности, адаптацией модели к новым поведенческим данным и повышением точности рекомендаций на основе выбранных продуктов без радикальной перестройки уже созданной архитектуры."
    )


def main() -> None:
    doc = Document(WORK_PATH)
    fix_paragraphs_and_captions(doc)
    replace_references(doc)
    doc.save(FINAL_PATH)

    # Keep only the final working DOCX requested by the user.
    for extra in [
        ROOT / "Диплом Абдулин Дмитрий АИС-42.docx",
        ROOT / "Диплом Абдулин Дмитрий АИС-42 - backup.updated.docx",
        ROOT / "tmp_backup_compare.docx",
    ]:
        if extra.exists():
            extra.unlink()

    print(f"Finalized: {FINAL_PATH}")


if __name__ == "__main__":
    main()
