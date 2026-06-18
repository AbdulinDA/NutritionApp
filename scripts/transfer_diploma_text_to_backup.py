from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(r"C:\Projects\NutritionApp")
SOURCE_PATH = ROOT / "Диплом Абдулин Дмитрий АИС-42.docx"
TARGET_PATH = ROOT / "Диплом Абдулин Дмитрий АИС-42 - backup.docx"
TEMP_TARGET_PATH = ROOT / "Диплом Абдулин Дмитрий АИС-42 - backup.updated.docx"

# The user asked not to touch the first 4 pages with administrative documents.
# In this template the main content starts from the "РЕФЕРАТ" page.
FIRST_EDITABLE_PARAGRAPH = 114
FIRST_EDITABLE_TABLE = 1


def set_paragraph_text_keep_format(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text_keep_format(cell, text: str) -> None:
    if cell.paragraphs:
        set_paragraph_text_keep_format(cell.paragraphs[0], text)
        for extra in cell.paragraphs[1:]:
            set_paragraph_text_keep_format(extra, "")
    else:
        cell.text = text


def transfer_paragraphs(source: Document, target: Document) -> None:
    for idx in range(FIRST_EDITABLE_PARAGRAPH, min(len(source.paragraphs), len(target.paragraphs))):
        set_paragraph_text_keep_format(target.paragraphs[idx], source.paragraphs[idx].text)


def transfer_tables(source: Document, target: Document) -> None:
    for table_idx in range(FIRST_EDITABLE_TABLE, min(len(source.tables), len(target.tables))):
        source_table = source.tables[table_idx]
        target_table = target.tables[table_idx]
        for row_idx in range(min(len(source_table.rows), len(target_table.rows))):
            source_row = source_table.rows[row_idx]
            target_row = target_table.rows[row_idx]
            for cell_idx in range(min(len(source_row.cells), len(target_row.cells))):
                set_cell_text_keep_format(target_row.cells[cell_idx], source_row.cells[cell_idx].text)


def main() -> None:
    source = Document(SOURCE_PATH)
    target = Document(TARGET_PATH)

    transfer_paragraphs(source, target)
    transfer_tables(source, target)

    try:
        target.save(TARGET_PATH)
        print(f"Updated backup file: {TARGET_PATH}")
    except PermissionError:
        target.save(TEMP_TARGET_PATH)
        print(f"Target is locked. Saved updated copy to: {TEMP_TARGET_PATH}")


if __name__ == "__main__":
    main()
